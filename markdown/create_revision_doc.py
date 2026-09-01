import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Helper function for setting cell background color
    def set_cell_background(cell, fill_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
        tcPr.append(shd)

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run('DAFTAR RENCANA UPDATE WEBSITE PEDIOCARE')
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(17)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 44, 89) # Navy Blue

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run('Hasil Analisis & Pencocokan Web Saat Ini terhadap Dokumen "Web Pediocare Revisi 2.docx"')
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    # Metadata Box
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ('Dokumen Sumber Acuan:', 'Web Pediocare Revisi 2.docx (Revisi Kedua Klien)'),
        ('Status Web Saat Ini:', 'Selesai Sinkronisasi Database & Aset Online, Menunggu Penerapan Revisi 2'),
        ('Tujuan Dokumen:', 'Memetakan seluruh poin perubahan, status kondisi saat ini, dan rencana tindakan teknis.')
    ]
    for i, (k, v) in enumerate(meta_data):
        r = meta_table.rows[i]
        r.cells[0].width = Inches(2.1)
        r.cells[1].width = Inches(4.4)
        set_cell_background(r.cells[0], 'F1F5F9')
        set_cell_background(r.cells[1], 'F8FAFC')
        
        p0 = r.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        p0.paragraph_format.space_before = Pt(2)
        run0 = p0.add_run(k)
        run0.font.bold = True
        run0.font.size = Pt(9.5)
        run0.font.color.rgb = RGBColor(30, 41, 59)
        
        p1 = r.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(2)
        run1 = p1.add_run(v)
        run1.font.size = Pt(9.5)
        run1.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 1: Ringkasan Eksekutif
    h1 = doc.add_heading('1. RINGKASAN EKSEKUTIF PERUBAHAN UTAMA', level=1)
    h1.style.font.color.rgb = RGBColor(15, 44, 89)
    h1.style.font.name = 'Arial'

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(8)
    p_intro.paragraph_format.line_spacing = 1.15
    p_intro.add_run(
        'Berdasarkan penelaahan mendalam terhadap file "Web Pediocare Revisi 2.docx", terdapat '
        '10 kelompok perubahan utama yang mencakup pergantian identitas warna dasar (dari Hijau menjadi Biru Tua / Navy kombinasi Biru Muda), '
        'penyempurnaan copywriting standar Kementerian Kesehatan Republik Indonesia, penambahan komponen interaktif photo auto-slider di halaman layanan medis, '
        'perapian data katalog produk brace & kaki palsu, pembersihan nama lama "Orthocare" menjadi "pediOcare", '
        'serta penambahan modul admin CMS untuk mengelola halaman "Alur Pasien" dan "Tentang Kami".'
    )

    # Heading 2: Tabel Rekapitulasi
    h2 = doc.add_heading('2. TABEL REKAPITULASI PENCOCOKAN DOKUMEN REVISI 2 VS KONDISI WEB SAAT INI', level=1)
    h2.style.font.color.rgb = RGBColor(15, 44, 89)
    h2.style.font.name = 'Arial'

    rekap_table = doc.add_table(rows=1, cols=4)
    rekap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['No', 'Item / Komponen', 'Status Saat Ini di Web', 'Rencana Tindakan Update (Revisi 2)']
    hdr_cells = rekap_table.rows[0].cells
    hdr_widths = [Inches(0.4), Inches(1.8), Inches(2.1), Inches(2.2)]

    for idx, heading in enumerate(headers):
        hdr_cells[idx].width = hdr_widths[idx]
        set_cell_background(hdr_cells[idx], '0F2C59')
        p = hdr_cells[idx].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(heading)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)

    table_rows_data = [
        ('1', 'Tema Warna Utama', 'Masih dominan Hijau (#306D29) pada tombol, badge, hero wave, navbar, & ikon.', 'Ganti total seluruh skema warna hijau menjadi Biru Tua (Navy #0F2C59 / #1E3A8A) gradasi Biru Muda (#38BDF8 / #E0F2FE).'),
        ('2', 'Nama Brand & Tagline', 'Tertulis pediOcare, tapi beberapa teks masih ada koma atau variasi spasi.', 'Standarkan nama: "pediOcare" dan tagline tepat di bawahnya tanpa koma: "Care your milestone".'),
        ('3', 'Poin Keunggulan (USP)', 'Masih teks lama (Teknologi 3D Scanning & Custom fitting presisi).', 'Ubah Poin 1 menjadi: "Teknologi terkini dengan standar pelayanan & alat customize yang presisi". Poin 2 & 3 disesuaikan.'),
        ('4', 'Section Layanan & Ikon', 'Judul: Layanan Unggulan Kami. Ikon orang kecil biasa.', 'Ubah susunan judul: "Layanan Orthosis Prosthesis Dan Alat Bantu Ortopedi". Ganti ikon orang kecil menjadi logo disabilitas (kursi roda).'),
        ('5', 'Inovasi Fabrikasi Modern', 'Belum ada kalimat pengantar baru dan masih memakai foto lama.', 'Tambahkan kalimat pengantar: "pediOcare akan terus berkembang menuju inovasi fabrikasi modern" dan ganti foto dengan image12.png.'),
        ('6', 'Layanan Prostesis', 'Teks umum Prosthetics, belum ada auto-slider di samping deskripsi.', 'Ganti judul menjadi "Prostesis (Kaki dan tangan Tiruan)", pasang copy Kemenkes RI, buat photo auto-slider di samping teks & box konsultasi di bawah.'),
        ('7', 'Layanan Bracing & Support', 'Belum ada auto-slider foto samping deskripsi.', 'Tambahkan photo slider otomatis di samping deskripsi layanan (image9.png) & box konsultasi medis di bawahnya.'),
        ('8', 'Layanan Skoliosis', 'Teks deskripsi umum, belum ada auto-slider khusus.', 'Update copywriting teknis (metode non-operatif POP, Cheneau 3-point pressure) & pasang auto-slider foto skoliosis (image1.png).'),
        ('9', 'Katalog & Card Produk', 'Nama produk: Korset Skoliosis TLSO, Kaki Palsu Bawah Lutut.', 'Ubah nama & deskripsi: "Skoliosis Brace", "Kaki Palsu Custom Made", "Skoliosis Brace with 3D Correction". Warna card biru tua.'),
        ('10', 'Aset Foto & Watermark', 'Beberapa foto/slide masih berlogo/teks "Orthocare".', 'Ganti seluruh materi foto & slide yang berlogo "Orthocare" dengan aset bertuliskan/berlogo "pediOcare".'),
        ('11', 'Menu Admin: Alur Pasien & Tentang Kami', 'Belum tersedia menu/tab khusus di CMS admin untuk mengedit 9 Tahapan Alur Pasien dan halaman Tentang Kami.', 'Tambahkan modul CMS Admin baru: Menu "Alur Pasien" (9 tahapan klinis) dan Menu/Tab "Tentang Kami" (Visi, Misi, Sejarah, Galeri) agar dapat diedit dinamis tanpa utak-atik code.')
    ]

    for row_data in table_rows_data:
        row = rekap_table.add_row()
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = hdr_widths[c_idx]
            if int(row_data[0]) % 2 == 0:
                set_cell_background(cell, 'F8FAFC')
            else:
                set_cell_background(cell, 'FFFFFF')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            if c_idx == 0:
                r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 3: Rincian Teknis
    h3 = doc.add_heading('3. RINCIAN DETAIL DAN SPESIFIKASI UPDATE PER MODUL', level=1)
    h3.style.font.color.rgb = RGBColor(15, 44, 89)
    h3.style.font.name = 'Arial'

    sections_detail = [
        (
            'Modul 1: Transformasi Identitas Visual & Skema Warna Global',
            [
                'Warna Utama (Primary): Ganti dari Hijau (#306D29) menjadi Biru Tua Medis (Deep Navy #0F2C59 / Indigo-Blue #1E3A8A).',
                'Warna Aksen & Gradasi (Secondary/Light): Biru Muda / Sky Blue (#38BDF8 / #E0F2FE) untuk efek gradasi bagian dalam.',
                'Komponen Terdampak: Tailwind config di layouts/app.blade.php, SVG wave transitions, WhatsApp CTA buttons, badge status, background section (#f0fdfa -> #f0f7ff), card borders, dan hover states di seluruh halaman.'
            ]
        ),
        (
            'Modul 2: Header, Logo, & Hero Section',
            [
                'Nama Klinik: Pastikan penulisan baku adalah "pediOcare" (huruf O kapital).',
                'Tagline Klinik: "Care your milestone" diletakkan tepat di bawah nama klinik tanpa tanda koma.',
                'Poin Keunggulan Utama (V-Checklist):',
                '  • [V1] "Teknologi terkini dengan standar pelayanan & alat customize yang presisi."',
                '  • [V2] "Praktisi Ortotis Prostetis legal memiliki STR & Surat Ijin Praktik Dinas Kesehatan."',
                '  • [V3] "Pelayanan komprehensif dan paripurna (konsultasi gratis)."'
            ]
        ),
        (
            'Modul 3: Section Layanan Orthosis Prosthesis & Alat Bantu Ortopedi',
            [
                'Judul Section: Format ulang susunan teks menjadi:\n  "Layanan Orthosis Prosthesis Dan Alat Bantu Ortopedi"',
                'Ikon Layanan: Ganti ikon orang kecil dengan simbol disabilitas resmi (Kursi Roda / Wheelchair accessible icon).',
                'Styling Kartu: Warna dasar kartu putih bersih dengan aksen garis dan ikon bernuansa Biru Tua gradasi Biru Muda.'
            ]
        ),
        (
            'Modul 4: Section Inovasi Fabrikasi Modern',
            [
                'Teks Pengantar: Tambahkan kalimat sebelum judul "Inovasi Fabrikasi Modern":\n  "pediOcare akan terus berkembang menuju inovasi fabrikasi modern"',
                'Foto Workshop / Fabrikasi: Ganti gambar lama dengan file foto baru berkualitas tinggi yang dilampirkan pada revisi (image12.png).'
            ]
        ),
        (
            'Modul 5: Halaman Layanan Prostesis (Kaki & Tangan Tiruan)',
            [
                'Judul Layanan: Ganti menjadi "Prostesis (Kaki dan tangan Tiruan)".',
                'Deskripsi 1: "Pembuatan kaki dan tangan palsu sesuai dengan standar pelayanan yang ditetapkan oleh Kementrian Kesehatan Republik Indonesia".',
                'Deskripsi 2: "Solusi mengembalikan fungsi anggota gerak tubuh yang hilang karena amputasi/bawaan sejak lahir (restoration of Function)".',
                'Interaktif: Tambahkan Photo Auto-Slider (carousel bergerak otomatis) di samping kolom deskripsi.',
                'Penempatan Box: Posisikan box banner "Konsultasi Medis" tepat di bawah deskripsi utama.'
            ]
        ),
        (
            'Modul 6: Halaman Layanan Bracing & Orthopaedic Supports',
            [
                'Interaktif: Pasang komponen Photo Auto-Slider di samping deskripsi produk penyangga ortopedi (menggunakan foto baru image9.png dan galeri terkait).',
                'Penempatan Box: Letakkan box informasi konsultasi medis klinis langsung di bawah deskripsi.'
            ]
        ),
        (
            'Modul 7: Halaman Layanan Penanganan Skoliosis (Scoliosis Center)',
            [
                'Interaktif: Tambahkan Photo Auto-Slider di samping teks penjelasan skoliosis (image1.png dan galeri foto koreksi skoliosis).',
                'Copywriting Standar Medis:',
                '  • "Penanganan skoliosis dengan metode non operatif dengan mencetak langsung dengan POP, kemudian di-design Cheneau dengan fokus 3 point penekanan di sepanjang tulang belakang."',
                '  • "Melalui rektifikasi menyeluruh dan analisa rotasi tulang belakang berbasis radiologi, kami membuat brace skoliosis custom yang bekerja 3 point pressure pada lengkungan tulang belakang."'
            ]
        ),
        (
            'Modul 8: Katalog & Kartu Produk (Product Cards)',
            [
                'Produk Korset Skoliosis: Ganti nama menjadi "Skoliosis Brace" dengan deskripsi:\n  "Brace yang mengoreksi skoliosis dengan sistem tiga dimensi secara efektif mengoreksi dan menghambat bertambahnya kurva derajat skoliosis."',
                'Produk Kaki Palsu: Ganti nama menjadi "Kaki Palsu Custom Made" dengan deskripsi:\n  "Prostesis atas lutut & bawah lutut dengan mengutamakan kenyamanan pengguna dengan material berkualitas dan sesuai standar."',
                'Produk Koreksi 3D: Ganti nama menjadi "Skoliosis Brace with 3D Correction".',
                'Warna Badge & Tombol: Ganti semua badge "Ready Stock" dan tombol katalog menjadi kombinasi Biru Tua & Biru Muda.'
            ]
        ),
        (
            'Modul 9: Pembersihan Aset Foto & Branding (Orthocare -> pediOcare)',
            [
                'Audit Visual: Periksa seluruh banner, thumbnail, dan slide foto yang masih memuat watermark, teks, atau logo lama "Orthocare".',
                'Penggantian Aset: Ganti seluruhnya dengan aset foto bertuliskan/berlogo resmi "pediOcare".'
            ]
        ),
        (
            'Modul 10: Pengembangan Modul CMS Admin (Alur Pasien & Tentang Kami)',
            [
                'Kebutuhan: Klien menyampaikan bahwa pada menu admin saat ini belum ada modul untuk mengedit Alur Pasien dan halaman Tentang Kami.',
                'Fitur yang Akan Ditambahkan di Panel Admin (/admin):',
                '  1. Modul CMS "Alur Pasien": Form kelola 9 tahapan alur klinis (Pemeriksaan, Diagnosis, Pengukuran, Pencetakan, Rektifikasi, Fabrikasi, Pengepasan, Penyerahan, Evaluasi) lengkap dengan icon, judul, dan deskripsi.',
                '  2. Modul CMS "Tentang Kami": Form kelola visi, misi, deskripsi profil, nilai-nilai utama, dan upload multi-foto galeri kegiatan klinik secara fleksibel dari dashboard admin.'
            ]
        )
    ]

    for title, points in sections_detail:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(8)
        p_sec.paragraph_format.space_after = Pt(3)
        r_sec = p_sec.add_run(title)
        r_sec.font.bold = True
        r_sec.font.size = Pt(11)
        r_sec.font.color.rgb = RGBColor(15, 44, 89)
        
        for pt in points:
            p_pt = doc.add_paragraph(style='List Bullet')
            p_pt.paragraph_format.space_before = Pt(1)
            p_pt.paragraph_format.space_after = Pt(2)
            p_pt.paragraph_format.line_spacing = 1.15
            r_pt = p_pt.add_run(pt)
            r_pt.font.size = Pt(9.5)
            r_pt.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 4: Roadmap Eksekusi
    h4 = doc.add_heading('4. ROADMAP & TAHAPAN PENGERJAAN', level=1)
    h4.style.font.color.rgb = RGBColor(15, 44, 89)
    h4.style.font.name = 'Arial'

    roadmap_p = doc.add_paragraph()
    roadmap_p.paragraph_format.line_spacing = 1.15
    roadmap_p.add_run(
        'Rencana pengerjaan akan dilakukan secara bertahap dan terstruktur sebagai berikut:\n'
        '• Tahap 1: Pengubahan Global Theme & Color Palette (CSS Tailwind, Wave SVG, Tombol, Background, Badge) dari Hijau ke Biru Tua & Biru Muda.\n'
        '• Tahap 2: Update Copywriting, Tagline, dan USP pada Header, Hero Section, dan Footer.\n'
        '• Tahap 3: Implementasi Photo Auto-Slider dan Layout Box Konsultasi Medis pada halaman Layanan Prostesis, Bracing, dan Skoliosis.\n'
        '• Tahap 4: Update Data Produk (Skoliosis Brace, Kaki Palsu Custom Made, 3D Correction) dan Pembersihan Aset Foto Orthocare -> pediOcare.\n'
        '• Tahap 5: Pembuatan Modul Admin CMS baru untuk Alur Pasien (9 Tahapan) & Halaman Tentang Kami agar dapat dikelola langsung oleh admin.\n'
        '• Tahap 6: Verifikasi tampilan responsif (Desktop & Mobile) dan pengujian menyeluruh.'
    )

    output_path = r'D:\var\www\html\Laravel\ortotik\markdown\Daftar_Rencana_Update_Revisi_2_pediOcare.docx'
    doc.save(output_path)
    print('Successfully generated docx at:', output_path)

if __name__ == '__main__':
    create_document()
